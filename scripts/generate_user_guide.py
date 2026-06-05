"""
Generates the Wizdent CRM navigation + functionality guide as a .docx file.
Run:  python3 scripts/generate_user_guide.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- palette ----------
NAVY = RGBColor(0x0F, 0x17, 0x2A)
BLUE = RGBColor(0x25, 0x63, 0xEB)
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
SLATE = RGBColor(0x47, 0x55, 0x69)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x05, 0x96, 0x69)
RED = RGBColor(0xDC, 0x26, 0x26)

HEADER_FILL = "1F2937"   # dark slate for table header rows
BAND_FILL = "EEF2FF"     # light indigo band
CRED_FILL = "ECFDF5"     # light green for credential rows

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=9.5, align="left", italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p

def shade_paragraph(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "2563EB")
        pbdr.append(bottom)
        pPr.append(pbdr)
        p.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = BLUE
        p.space_after = Pt(2)
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = INDIGO
    return p

def add_body(text, bullet=False, size=10.5, italic=False, color=None, bold=False):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)
    return p

def add_steps(steps):
    for s in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(s).font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(2)

def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

def header_row(table, labels):
    cells = table.rows[0].cells
    for i, lab in enumerate(labels):
        set_cell_bg(cells[i], HEADER_FILL)
        set_cell_text(cells[i], lab, bold=True, color=WHITE, size=9.5,
                      align="center" if i else "left")

# =====================================================================
# COVER
# =====================================================================
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("WIZDENT")
r.bold = True
r.font.size = Pt(40)
r.font.color.rgb = BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Field Sales Automation & Visit Management CRM")
r.font.size = Pt(14)
r.font.color.rgb = SLATE

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("User Navigation & Functionality Guide")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = NAVY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Walkthrough from Login → Every Feature  •  Role-wise Unique Capabilities  •  Test Login Credentials")
r.italic = True
r.font.size = Pt(10.5)
r.font.color.rgb = SLATE

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta2.add_run("Document generated 5 June 2026  •  Web application (React + Firebase)")
r.font.size = Pt(9)
r.font.color.rgb = SLATE

doc.add_paragraph()

# =====================================================================
# TABLE OF CONTENTS
# =====================================================================
add_heading("Contents", 1)
toc = [
    "1.  About the Application",
    "2.  Logging In (Step-by-Step)",
    "3.  Login Credentials for Every User in the Database",
    "4.  User Roles & Access Matrix",
    "5.  The Sidebar — Global Navigation Map",
    "6.  Feature-by-Feature Walkthrough",
    "7.  Role-Wise Unique Features",
    "8.  Core Workflow — The Visit Lifecycle",
    "9.  Notes, Tips & Troubleshooting",
]
for t in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(t)
    run.font.size = Pt(10.5)
    run.bold = True
    run.font.color.rgb = SLATE

doc.add_page_break()

# =====================================================================
# 1. ABOUT
# =====================================================================
add_heading("1.  About the Application", 1)
add_body("Wizdent SFA CRM is a Sales-Force-Automation tool for a dental-products distribution business. "
         "Field sales representatives (Dealers) plan and execute visits to dental clinics, log products sold, "
         "demonstrated and sampled, and capture GPS-verified check-in / check-out. Managers monitor their "
         "territory and Administrators run the whole system, manage the team, the catalog and master data.")
add_body("Technology: a single-page React web app backed by Google Firebase (Authentication + Firestore "
         "database). It is a website — open it in any modern browser; it is also mobile-responsive.")
# Hosting ownership callout
note = doc.add_paragraph()
shade_paragraph(note, "FEF3C7")  # amber band
r = note.add_run("  ⚠  Hosting & ownership: ")
r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
r = note.add_run("The current live deployment runs on the project owner's PERSONAL server / Firebase "
                 "account and is for demo and evaluation only. For real production use, download (clone) "
                 "the repository locally and create a NEW deployment on the COMPANY's own servers — a "
                 "company-owned Firebase project, Hosting site, credentials and billing. Do not treat the "
                 "personal deployment as production.")
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
note.paragraph_format.space_after = Pt(8)

add_body("Three account types (roles) exist, each unlocking a different set of screens and powers:")
for txt, col in [("ADMIN — full system control", RED),
                 ("MANAGER — territory oversight & analytics", INDIGO),
                 ("DEALER — the field sales rep who actually runs visits", BLUE)]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(txt); run.font.size = Pt(10.5); run.bold = True; run.font.color.rgb = col

# =====================================================================
# 2. LOGGING IN
# =====================================================================
add_heading("2.  Logging In (Step-by-Step)", 1)
add_body("When you open the app you land on the Login screen (a dark page with the Wizdent “W” logo). "
         "There are two sign-in tabs: Google and Credentials.")

add_heading("Option A — Email / Password (Credentials tab)", 3)
add_steps([
    "Open the website. You are redirected to the Login page.",
    "Click the “Credentials” tab (toggle at the top of the login card).",
    "In the “Username / Email” field type the user’s email address (e.g. kislay@gmail.com). "
    "Tip: for the dedicated system admin you can simply type admin — it is auto-mapped to admin@wizdent.system.",
    "In the “Password” field type the password (12345678 for the standard test accounts — see Section 3).",
    "Click “Sign In”. On success you are taken to the Command Center dashboard.",
])

add_heading("Option B — Continue with Google", 3)
add_steps([
    "On the Login page keep the “Google” tab selected (it is the default).",
    "Click “Continue with Google” and pick your Google account in the popup.",
    "First-time Google users are auto-provisioned; the owner email parvvirmani07@gmail.com is auto-elevated to ADMIN.",
])
add_body("Logging out: open My Profile and click “Log Out”, or use the small log-out icon at the bottom of "
         "the sidebar next to your name.", italic=True, color=SLATE, size=9.5)

# =====================================================================
# 3. CREDENTIALS
# =====================================================================
add_heading("3.  Login Credentials for Every User in the Database", 1)
add_body("These are the actual user accounts currently provisioned in the Firestore “users” collection. "
         "The password for the standard accounts is the same for everyone:", bold=True)

pw = doc.add_paragraph()
pw.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pw.add_run("  Password (all standard accounts):   12345678  ")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = WHITE
shade_paragraph(pw, "059669")

add_body("")
data = [
    # role, name, email/username, password, sign-in, territory
    ("DEALER", "Apex Medical Distributors", "apex-medical-distributors@gmail.com", "12345678", "Email/Password", "MH01 – Mumbai (recommended demo dealer)"),
    ("DEALER", "Coromandel Meditech Solutions", "coromandel-meditech-solutions@gmail.com", "12345678", "Email/Password", "— (no territory set)"),
    ("DEALER", "Deccan Surgical Wholesalers", "deccan-surgical-wholesalers@gmail.com", "12345678", "Email/Password", "— (no territory set)"),
    ("DEALER", "Matrix Pharma & Surgicals", "matrix-pharma--surgicals@gmail.com", "12345678", "Email/Password", "— (no territory set)"),
    ("DEALER", "Vedic Healthcare Logistics", "vedic-healthcare-logistics@gmail.com", "12345678", "Email/Password", "— (no territory set)"),
    ("MANAGER", "Kislay Kritesh", "kislay@gmail.com", "12345678", "Email/Password", "MH01 – Mumbai - South"),
    ("ADMIN", "Parv Admin", "parv.admin@wizdent.com", "12345678", "Email/Password", "— (global access)"),
    ("ADMIN", "System Admin", "admin@wizdent.system  (or type: admin)", "(set by owner)", "Email/Password", "— (global access)"),
    ("ADMIN", "Parv Virmani (Owner)", "parvvirmani07@gmail.com", "(Google account)", "Google Sign-In", "— (global access)"),
]
table = doc.add_table(rows=1, cols=6)
style_table(table)
header_row(table, ["Role", "Name", "Email / Username", "Password", "Sign-in method", "Territory"])
for role, name, email, pwd, method, terr in data:
    cells = table.add_row().cells
    rc = {"ADMIN": RED, "MANAGER": INDIGO, "DEALER": BLUE}[role]
    set_cell_text(cells[0], role, bold=True, color=rc, size=9)
    set_cell_text(cells[1], name, size=9)
    set_cell_text(cells[2], email, size=8.5)
    set_cell_text(cells[3], pwd, bold=True, size=9,
                  color=GREEN if pwd == "12345678" else SLATE)
    set_cell_text(cells[4], method, size=8.5)
    set_cell_text(cells[5], terr, size=8.5)
    if pwd == "12345678":
        set_cell_bg(cells[3], CRED_FILL)

add_body("")
add_body("Important notes on the ADMIN accounts:", bold=True, color=RED)
add_body("Parv Admin (parv.admin@wizdent.com) is a standard email/password administrator and DOES sign in "
         "with 12345678 — use this for the simplest admin login.", bullet=True, size=9.5)
add_body("Parv Virmani (parvvirmani07@gmail.com) is the system owner and signs in with Google, not a password. "
         "This email is hard-coded as a super-admin.", bullet=True, size=9.5)
add_body("System Admin (admin@wizdent.system) is the dedicated username-login admin — on the login screen you "
         "may simply type the username admin. Its password is managed by the owner and is NOT 12345678; ask the "
         "owner for it, or use the Parv Admin / Parv Virmani accounts instead.", bullet=True, size=9.5)
add_body("Verified: the Parv Admin account, the 1 Manager and all 5 Dealer accounts above log in successfully "
         "with 12345678.", bullet=True, size=9.5, italic=True, color=GREEN)
add_body("Tip: To explore the full Dealer experience (seeing clinics, planning visits) sign in as "
         "Apex Medical Distributors — it is the only dealer whose profile has a territory (MH01), which is what "
         "controls which clinics a dealer can see. An admin can assign territories to the other dealers from the "
         "Control Tower if needed.", bullet=True, size=9.5, color=SLATE)

doc.add_page_break()

# =====================================================================
# 4. ACCESS MATRIX
# =====================================================================
add_heading("4.  User Roles & Access Matrix", 1)
add_body("What each role can reach and do. ✓ = available, ✗ = blocked / hidden.")

matrix = [
    ("Command Center (Dashboard)", "All visits & KPIs", "Own territory only", "Own visits only"),
    ("My Profile", "✓", "✓", "✓ + personal KPIs & Dealer Code"),
    ("Account Management (Clinics/Dealers)", "View, Add, Bulk-Delete (all)", "View + Add (territory)", "View clinics in own territory"),
    ("Contact Management", "✓", "✓", "✓"),
    ("Team Directory (Users)", "✓ Full directory", "✗ Blocked (hidden)", "✓ Limited (territory)"),
    ("Analytics Reports", "✓ All tabs", "✓ Territory data", "✗ Restricted screen"),
    ("Control Tower / Identity Monitor (Admin)", "✓", "✗ Access denied", "✗ Access denied"),
    ("Visit Planning (list)", "All + delete visits", "Territory visits", "Own visits"),
    ("Plan a New Visit", "✓", "✗ Redirected", "✓"),
    ("Visit Check-in / Check-out (GPS)", "✗", "✗", "✓ (the field rep)"),
    ("Product Catalog", "View + Add + Delete SKUs", "View", "View"),
    ("Bulk Data Import (XLSX/CSV)", "✓ (inside Control Tower)", "✗", "✗"),
]
t = doc.add_table(rows=1, cols=4)
style_table(t)
header_row(t, ["Feature / Screen", "ADMIN", "MANAGER", "DEALER"])
for i, (feat, a, m, d) in enumerate(matrix):
    cells = t.add_row().cells
    set_cell_text(cells[0], feat, bold=True, size=9)
    set_cell_text(cells[1], a, size=8.5, align="center")
    set_cell_text(cells[2], m, size=8.5, align="center")
    set_cell_text(cells[3], d, size=8.5, align="center")
    if i % 2 == 1:
        for c in cells:
            set_cell_bg(c, "F8FAFC")

# =====================================================================
# 5. SIDEBAR
# =====================================================================
add_heading("5.  The Sidebar — Global Navigation Map", 1)
add_body("After login every screen shares the same dark left sidebar. On mobile, tap the ☰ menu icon in the "
         "top bar to open it. Items are grouped into three sections:")

add_heading("Operational Hub", 3)
for label, dest in [
    ("Command Center", "Dashboard with KPIs and charts (home / “/”)"),
    ("My Profile", "Your identity, role, territory and (for dealers) performance stats"),
    ("Account Management", "Clinics & dealer partners"),
    ("Contact Management", "Doctors / contacts at clinics"),
    ("Team Directory", "All personnel — hidden for Managers"),
    ("Analytics Reports", "Dealer / clinic / margin reports — Managers & Admins"),
]:
    p = doc.add_paragraph(style="List Bullet")
    rb = p.add_run(label + " — "); rb.bold = True; rb.font.size = Pt(10)
    rt = p.add_run(dest); rt.font.size = Pt(10); rt.font.color.rgb = SLATE

add_heading("System Admin  (visible to Admins only)", 3)
p = doc.add_paragraph(style="List Bullet")
rb = p.add_run("Identity Monitor — "); rb.bold = True; rb.font.size = Pt(10)
rt = p.add_run("the Control Tower: create users, provision dealers, bulk import."); rt.font.size = Pt(10); rt.font.color.rgb = SLATE

add_heading("Supply Chain", 3)
for label, dest in [
    ("Visit Planning", "List of visits + “Plan New Visit” (dealers)"),
    ("Product Catalog", "Dental product / SKU master list"),
]:
    p = doc.add_paragraph(style="List Bullet")
    rb = p.add_run(label + " — "); rb.bold = True; rb.font.size = Pt(10)
    rt = p.add_run(dest); rt.font.size = Pt(10); rt.font.color.rgb = SLATE

add_body("The footer of the sidebar shows your initials, name and role, plus a quick Log-Out button. "
         "The top bar shows breadcrumbs, a global search box and your signed-in email.", italic=True, color=SLATE, size=9.5)

doc.add_page_break()

# =====================================================================
# 6. FEATURE WALKTHROUGH
# =====================================================================
add_heading("6.  Feature-by-Feature Walkthrough", 1)

def feature(title, route, paras):
    add_heading(f"{title}", 2)
    rp = doc.add_paragraph()
    r = rp.add_run(f"Route: {route}")
    r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = SLATE
    rp.paragraph_format.space_after = Pt(2)
    for para in paras:
        add_body(para, bullet=True, size=10)

feature("6.1  Command Center (Dashboard)", "/", [
    "Landing screen after login. Four KPI cards: Total Revenue (₹), Field Visits, Conversion Rate %, Active Clinics.",
    "Two charts: a 7-day revenue area chart and a “visits per day” bar chart, plus average productivity and growth index.",
    "Data is scoped to your role: Admin sees everything, Manager sees their territory, Dealer sees only their own visits.",
])
feature("6.2  My Profile", "/profile", [
    "Shows your avatar (initials), display name, role badge, email, territory and join date.",
    "Dealers additionally get three KPI tiles (Total Visits, Total Revenue, Conversions %) and a “Recent Visits Logged” list.",
    "Displays your “Account Credentials” card — permission tier, assigned territory and (dealers) the official Dealer Code.",
    "Contains a Log Out button.",
])
feature("6.3  Account Management", "/accounts", [
    "Lists Clinics and Dealer partners in a filterable table (All / Clinics / Dealers tabs) with search.",
    "“Add Account” opens a modal to register a Clinic (type, territory, clinic type, avg patients, address) or a "
    "Dealer (dealer code, reseller landing cost). Dealers can only add Clinics, not Dealer accounts.",
    "Admins get checkboxes + “Bulk Delete”, which also cascades to delete the account’s contacts, visits and sub-records.",
    "Click “Detail” on any row to open the Account Detail page.",
])
feature("6.4  Contact Management", "/contacts", [
    "Directory of doctors / contacts attached to clinics (name, specialty, mobile, email).",
    "Open a contact to view the Contact Detail page; contacts are also created on-the-fly while planning a visit.",
])
feature("6.5  Team Directory", "/users", [
    "Card grid of all personnel with role badge, email, territory and dealer code; search + filter by role.",
    "Admins see the full directory; Dealers see a territory-limited list.",
    "Managers are deliberately blocked here — they see a “Security Access Warning” screen, and the menu item is hidden for them.",
    "Click a card to open that user’s detail page.",
])
feature("6.6  Analytics Reports", "/reports", [
    "Tabbed analytics: Dealers, Clinics, Products and Margin Analysis, with an Export Datatable button.",
    "Dealers tab: visits logged, total revenue, conversion %, slab per dealer. Clinics tab: interactions, lifetime value, "
    "engagement status. Margin tab: estimated margin and margin % per dealer.",
    "Managers see only their territory’s figures; Admins see global. Dealers are shown a “Restricted Access” screen.",
])
feature("6.7  Control Tower / Identity Monitor (Admin only)", "/admin", [
    "Manual Onboarding form — create a single user: legal name, work email, security key (password), access tier "
    "(Manager / Dealer / Admin), geographic assignment, and dealer code for dealers.",
    "“Provision Existing Dealers” — one click auto-creates login accounts for every Dealer-type account "
    "(email = slugified-name@gmail.com, password = 12345678) and links them; results stream in a live feed.",
    "Bulk Data Import — upload an XLSX/CSV of products, accounts or contacts; columns are auto-mapped and written in batches.",
    "There is also a “Create Manager” flow with hierarchical Region → State → Area selection (RSM / State Manager / ASM).",
    "Non-admins who reach this URL get an “Access Denied” message.",
])
feature("6.8  Visit Planning (list)", "/visits", [
    "Table of visits: date/time, clinic, visit type, revenue, execution status (PLANNED / IN_PROGRESS / COMPLETED).",
    "Only Dealers see the “Plan New Visit” button. Only Admins see a delete (trash) action per visit.",
    "Click the chevron to open a visit’s detail page. Scope follows the role (own / territory / all).",
])
feature("6.9  Plan a New Visit", "/visits/new", [
    "Wizard to schedule a field visit: pick the Target Clinic and Primary Contact (searchable lookups), add extra "
    "attendees, a subject/objective, planned date, interaction type (Sell / Demo / Follow-up / Sampling / Promotion…) "
    "and activity protocol (Meeting / Workshop / Conference).",
    "“Quick Add Doctor” lets you create a new contact inline. System parameters: package SKU, currency, financial year.",
    "Restricted to Admins and Dealers — Managers are redirected back to the visit list.",
])
feature("6.10  Visit Detail & Execution", "/visits/:id", [
    "The operational heart of the app. Shows clinic, fulfillment dealer, primary contact, attendees, technical meta and notes.",
    "Dealers run the visit lifecycle here: “Check-in Now” captures GPS + timestamp (status → IN_PROGRESS); while in "
    "progress they add Products Sold (qty/price), Clinical Demos, Free Samples and Action Points; “Check-out & Finish” "
    "captures GPS again, computes duration, total revenue, PIL sales value and marks the visit COMPLETED / converted.",
    "Hero tiles show Net Revenue, PIL Sales Value, Visit Type, Duration/Status and Visit Date. “Print Report” is available to all.",
])
feature("6.11  Product Catalog", "/products", [
    "Real-time card grid of dental SKUs: name, division, SKU code, brand, three price tiers (Standard / Dealer / Dentist), "
    "pack size and active status.",
    "Everyone can browse. Only Admins can “Add SKU” (modal) or select + bulk-delete products.",
    "Click a card to open the Product Detail page.",
])

doc.add_page_break()

# =====================================================================
# 7. ROLE-WISE UNIQUE FEATURES
# =====================================================================
add_heading("7.  Role-Wise Unique Features", 1)
add_body("The capabilities below are the things only that role can do (or experience differently).")

def role_block(title, color, items):
    p = doc.add_paragraph()
    shade_paragraph(p, "1F2937")
    r = p.add_run("   " + title + "   ")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = WHITE
    for it in items:
        add_body(it, bullet=True, size=10)

role_block("ADMIN — unique powers", RED, [
    "Access to the Control Tower (/admin): the only role that can create new users of ANY tier.",
    "One-click “Provision Existing Dealers” to auto-mint dealer logins (password 12345678).",
    "Bulk Data Import of products, accounts and contacts from XLSX/CSV.",
    "View the complete Team Directory across all territories.",
    "Add and bulk-delete Products (SKUs); bulk-delete Accounts (cascading) and delete individual Visits.",
    "Global, unfiltered view of every dashboard, report, account and visit in the system.",
    "Create Manager accounts with hierarchical Region → State → Area assignment (RSM / State Manager / ASM).",
])
role_block("MANAGER — unique powers", INDIGO, [
    "Territory-scoped oversight: dashboard, accounts, visits and reports are all auto-filtered to the manager’s territory.",
    "Full Analytics Reports access (Dealers / Clinics / Products / Margin) — a screen Dealers cannot open.",
    "Can add and update accounts within their territory.",
    "Deliberately CANNOT open the Team Directory (compliance lock) and cannot Plan a new visit (oversight, not field-execution).",
])
role_block("DEALER — unique powers", BLUE, [
    "The only role that can Plan a New Visit and run it: GPS Check-in / Check-out on the Visit Detail screen.",
    "Logs the commercial outcome of a visit — Products Sold (qty & price), Clinical Demos, Free Samples and Action Points.",
    "Gets a personalised Profile with KPI tiles (Total Visits, Revenue, Conversions) and a Recent Visits feed.",
    "Sees their own official Dealer Code and a clinic list scoped to their assigned territory.",
    "Data is private to them — a dealer only ever sees their own visits.",
])

# =====================================================================
# 8. CORE WORKFLOW
# =====================================================================
add_heading("8.  Core Workflow — The Visit Lifecycle", 1)
add_body("The end-to-end happy path that ties the product together (performed by a Dealer):")
add_steps([
    "Sign in as a Dealer (e.g. apex-medical-distributors@gmail.com / 12345678).",
    "Go to Visit Planning → “Plan New Visit”. Choose the target clinic and a doctor (or Quick-Add one), set the "
    "date, interaction type and objective, then “Schedule Field Interaction”. The visit appears as PLANNED.",
    "On the visit day, open the visit and press “Check-in Now” — allow location access. Status becomes IN_PROGRESS.",
    "During the meeting, add Products Sold (quantity & price), any Clinical Demos, Free Samples and type Action Points.",
    "Press “Check-out & Finish” — location is captured again, duration and totals are computed, and the visit is "
    "marked COMPLETED (and converted if anything was sold).",
    "Revenue and conversion now roll up into the Dashboard, the dealer’s Profile, and the Manager/Admin Reports.",
])

# =====================================================================
# 9. NOTES / TIPS
# =====================================================================
add_heading("9.  Notes, Tips & Troubleshooting", 1)
add_body("Deployment ownership. The current live site is hosted on the owner's personal server / "
         "Firebase account for demo only. For real company use, clone the repo locally and deploy a fresh "
         "copy to the company's own Firebase project and servers (see the project README for steps).",
         bullet=True, size=10, color=RED)
add_body("Territory drives visibility. A user’s assigned territory determines which clinics, accounts and visits "
         "they can see. A dealer with no territory will see an empty clinic list — sign in as Apex Medical "
         "Distributors (MH01), or have an admin assign a territory.", bullet=True, size=10)
add_body("Territory codes you’ll encounter in the data: MH01 (Mumbai-South), DL01 (Delhi-North), KA01 "
         "(Bangalore-Central), TN01 (Chennai), AP01 (Hyderabad). The sample database holds 20 clinics and 5 dealer "
         "partners spread across these.", bullet=True, size=10)
add_body("Adding a user in Firestore is not enough to log in — the account must also exist in Firebase "
         "Authentication. The Control Tower’s create/provision flows handle both at once.", bullet=True, size=10)
add_body("Email/Password sign-in must be enabled in the Firebase console (Authentication → Sign-in method) for "
         "credential login to work.", bullet=True, size=10)
add_body("Check-in / Check-out needs the browser’s location permission. If you deny it, the action fails with an error.",
         bullet=True, size=10)
add_body("Forgot which screen you’re on? The breadcrumb in the top bar always shows the current section.",
         bullet=True, size=10)

# closing line
doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run("— End of Guide —")
r.italic = True; r.font.color.rgb = SLATE; r.font.size = Pt(10)

out = "Wizdent_CRM_User_Guide.docx"
doc.save(out)
print("Saved", out)
